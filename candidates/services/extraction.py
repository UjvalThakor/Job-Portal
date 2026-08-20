"""
Resume -> raw text extraction pipeline.

Flow (mirrors the architecture diagram):
    1. Try native text extraction (fast path, works for digitally-generated
       PDFs / DOCX).
    2. If that yields little/no text (i.e. a scanned resume or an image
       upload), fall back to: PDF -> images (PyMuPDF) -> OpenCV
       preprocessing -> Tesseract OCR.
"""
import io
import logging

import numpy as np
import cv2
from PIL import Image

logger = logging.getLogger(__name__)

MIN_NATIVE_TEXT_CHARS = 200  # below this, we assume the PDF is scanned/image-based


class ExtractionResult:
    def __init__(self, text: str, ocr_used: bool, pages: int = 1):
        self.text = text
        self.ocr_used = ocr_used
        self.pages = pages


# ---------------------------------------------------------------------------
# Native extraction (fast path)
# ---------------------------------------------------------------------------

def extract_pdf_native_text(file_bytes: bytes) -> str:
    import pdfplumber
    text_chunks = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            text_chunks.append(page_text)
    return "\n".join(text_chunks).strip()


def extract_docx_text(file_bytes: bytes) -> str:
    import docx
    doc = docx.Document(io.BytesIO(file_bytes))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts).strip()


# ---------------------------------------------------------------------------
# PDF -> images (for the OCR fallback path)
# ---------------------------------------------------------------------------

def pdf_to_images(file_bytes: bytes, dpi: int = 300):
    """High-quality PDF page rasterization using PyMuPDF."""
    import fitz  # PyMuPDF
    images = []
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    zoom = dpi / 72
    matrix = fitz.Matrix(zoom, zoom)
    for page in doc:
        pix = page.get_pixmap(matrix=matrix)
        img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
        images.append(img)
    doc.close()
    return images


# ---------------------------------------------------------------------------
# OpenCV preprocessing
# ---------------------------------------------------------------------------

def preprocess_image(pil_image: Image.Image) -> np.ndarray:
    """Resize, denoise, sharpen, deskew, enhance contrast, adaptive-threshold.
    Returns a binarized numpy array ready for OCR."""
    img = cv2.cvtColor(np.array(pil_image.convert("RGB")), cv2.COLOR_RGB2BGR)

    # Resize (upscale small images to help OCR accuracy)
    h, w = img.shape[:2]
    if max(h, w) < 1800:
        scale = 1800 / max(h, w)
        img = cv2.resize(img, None, fx=scale, fy=scale, interpolation=cv2.INTER_CUBIC)

    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)

    # Denoise
    gray = cv2.fastNlMeansDenoising(gray, h=10)

    # Sharpen
    kernel = np.array([[0, -1, 0], [-1, 5, -1], [0, -1, 0]])
    gray = cv2.filter2D(gray, -1, kernel)

    # Contrast enhancement (CLAHE)
    clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
    gray = clahe.apply(gray)

    # Deskew
    gray = _deskew(gray)

    # Adaptive threshold (binarize)
    binarized = cv2.adaptiveThreshold(
        gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 31, 11
    )
    return binarized


def _deskew(gray: np.ndarray) -> np.ndarray:
    thresh = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY_INV | cv2.THRESH_OTSU)[1]
    coords = np.column_stack(np.where(thresh > 0))
    if coords.shape[0] < 20:
        return gray
    angle = cv2.minAreaRect(coords)[-1]
    if angle < -45:
        angle = -(90 + angle)
    else:
        angle = -angle
    if abs(angle) < 0.5:
        return gray
    (h, w) = gray.shape[:2]
    center = (w // 2, h // 2)
    M = cv2.getRotationMatrix2D(center, angle, 1.0)
    return cv2.warpAffine(gray, M, (w, h), flags=cv2.INTER_CUBIC, borderMode=cv2.BORDER_REPLICATE)


# ---------------------------------------------------------------------------
# OCR
# ---------------------------------------------------------------------------

def ocr_image(preprocessed: np.ndarray) -> str:
    import pytesseract
    return pytesseract.image_to_string(preprocessed)


def ocr_pil_image(pil_image: Image.Image) -> str:
    preprocessed = preprocess_image(pil_image)
    return ocr_image(preprocessed)


# ---------------------------------------------------------------------------
# Orchestration
# ---------------------------------------------------------------------------

def extract_text(file_bytes: bytes, ext: str) -> ExtractionResult:
    ext = ext.lower()

    if ext == '.docx':
        text = extract_docx_text(file_bytes)
        return ExtractionResult(text=text, ocr_used=False)

    if ext in ('.png', '.jpg', '.jpeg'):
        pil_image = Image.open(io.BytesIO(file_bytes))
        text = ocr_pil_image(pil_image)
        return ExtractionResult(text=text, ocr_used=True)

    if ext == '.pdf':
        native_text = ""
        try:
            native_text = extract_pdf_native_text(file_bytes)
        except Exception as e:
            logger.warning("Native PDF text extraction failed: %s", e)

        if len(native_text) >= MIN_NATIVE_TEXT_CHARS:
            return ExtractionResult(text=native_text, ocr_used=False)

        # Fallback: scanned / image-based PDF -> OCR pipeline
        logger.info("PDF appears to be scanned/image-based; falling back to OCR")
        images = pdf_to_images(file_bytes)
        ocr_chunks = [ocr_pil_image(img) for img in images]
        ocr_text = "\n".join(ocr_chunks).strip()

        # Use whichever produced more content
        final_text = ocr_text if len(ocr_text) > len(native_text) else native_text
        return ExtractionResult(text=final_text, ocr_used=True, pages=len(images))

    raise ValueError(f"Unsupported extension for extraction: {ext}")
